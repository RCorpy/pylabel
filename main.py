import sys
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QApplication, QMainWindow
from PyQt5.QtGui import QIcon

from docx import Document
from docx.shared import Inches, Pt



def window():
    app = QApplication(sys.argv)
    win = QMainWindow()
    win.setGeometry(200,200,1200, 700)
    win.setWindowTitle = "Label maker"
    win.setWindowIcon(QIcon("icon.jpg"))

    lbl_titulo = QtWidgets.QLabel(win)
    lbl_titulo.setText("Titulo: ")
    lbl_titulo.move(50,50)

    lbl_surtitulo = QtWidgets.QLabel(win)
    lbl_surtitulo.setText("Enter surtitulo: ")
    lbl_surtitulo.move(50,90)

    txt_titulo = QtWidgets.QLineEdit(win)
    txt_titulo.move(200,50)

    txt_surtitulo = QtWidgets.QLineEdit(win)
    txt_surtitulo.move(200,90)

    def clicked(self):
        print("button clicked")
        
        doc = Document()
        
        #p = doc.add_paragraph()
        #r = p.add_run()
        #r.add_text('Good Morning every body,This is my ')
        #r.add_picture('icon.jpg',width=Inches(4.0), height=Inches(.7))
        #r.add_text(' do you like it?')
        #p1 = doc.add_paragraph("First Paragraph.")
        #p2 = doc.add_paragraph("Second Paragraph.")

        table = doc.add_table(rows=3, cols=2)
        table.cell(0,0).text = "0,0"
        table.cell(0,1).text = "0,1"
        table.cell(1,0).text = "1,0"
        table.cell(1,1).text = "1,1"
        doc.save('tabletest.docx')

        

        
    def loadClicked():
        print("load clicked")
        #print(txt_titulo.text() + " " + txt_surtitulo.text())
        document = Document("og.docx")
        #document.paragraphs[0].text = "YEAH"
        #print(document.paragraphs[0].text)
        #fullText = []
        #for para in document.paragraphs:
        #    print("hi")
        #    fullText.append(para.text)
        #print('\n'.join(fullText))
        table = document.tables[0]

        titlerun = table.cell(1,0).add_paragraph().add_run()
        font = titlerun.font
        font.name = "Arial Black"
        font.size = Pt(18)
        titlerun.add_text(txt_titulo.text())

        # second part
        titlerun2 = table.cell(1,0).add_paragraph().add_run()
        font2 = titlerun2.font
        font2.name = "Arial"
        font2.size = Pt(7)
        titlerun2.add_text(txt_titulo.text())
        print("done")

        document.save("out.docx")

    btn_save = QtWidgets.QPushButton(win)
    btn_save.move(200,180)
    btn_save.clicked.connect(clicked)
    btn_save.setText("Save")

    btn_load = QtWidgets.QPushButton(win)
    btn_load.move(200,130)
    btn_load.clicked.connect(loadClicked)
    btn_load.setText("Load")


    win.show()
    sys.exit(app.exec_())
    


window()